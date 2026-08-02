#!/usr/bin/env python3
"""Build a provenance fixture from public AKKE source documents.

The script downloads three representative originals (PDF, DOCX and XLSX),
extracts human-readable locators, and writes a deterministic JSON fixture for
the isolated D1 test database. It never reads or writes the production tables.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import urllib.request
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter


SOURCES = [
    {
        "id": "91c87aff-19c5-45fa-916c-74524430e0d0",
        "title": "装修工程避坑手册（2024年，100页）",
        "filename": "装修工程避坑手册（2024年，100页）.pdf",
        "local_name": "renovation-pitfalls-2024.pdf",
        "category": "避坑手册",
        "file_type": "pdf",
        "updated_at": "2026-07-20T09:31:56.993+00:00",
        "url": "https://pub-667af1183df84306ad1f8d1cec9bd192.r2.dev/docs/%E9%81%BF%E5%9D%91%E6%89%8B%E5%86%8C/%E8%A3%85%E4%BF%AE%E5%B7%A5%E7%A8%8B%E9%81%BF%E5%9D%91%E6%89%8B%E5%86%8C%EF%BC%882024%E5%B9%B4%EF%BC%8C100%E9%A1%B5%EF%BC%89.pdf",
    },
    {
        "id": "99f537aa-df08-4cff-8932-24a27ca122c8",
        "title": "装饰装修验收标准",
        "filename": "装饰装修验收标准.docx",
        "local_name": "decoration-acceptance-standard.docx",
        "category": "装修验收标准",
        "file_type": "docx",
        "updated_at": "2026-07-20T09:31:45.543+00:00",
        "url": "https://pub-667af1183df84306ad1f8d1cec9bd192.r2.dev/docs/%E9%AA%8C%E6%94%B6%E6%A0%87%E5%87%86/%E8%A3%85%E9%A5%B0%E8%A3%85%E4%BF%AE%E9%AA%8C%E6%94%B6%E6%A0%87%E5%87%86.docx",
    },
    {
        "id": "2080854c-0238-46a5-a9c1-a91c47c9398a",
        "title": "【表格】精装修验收标准",
        "filename": "【表格】精装修验收标准.xlsx",
        "local_name": "refined-decoration-acceptance.xlsx",
        "category": "装修验收标准",
        "file_type": "xlsx",
        "updated_at": "2026-07-20T09:28:32.706+00:00",
        "url": "https://pub-667af1183df84306ad1f8d1cec9bd192.r2.dev/docs/%E9%AA%8C%E6%94%B6%E6%A0%87%E5%87%86/%E3%80%90%E8%A1%A8%E6%A0%BC%E3%80%91%E7%B2%BE%E8%A3%85%E4%BF%AE%E9%AA%8C%E6%94%B6%E6%A0%87%E5%87%86.xlsx",
    },
]


def compact(text: str, limit: int | None = None) -> str:
    value = re.sub(r"\s+", " ", text or "").strip()
    if limit and len(value) > limit:
        return value[: limit - 1].rstrip() + "…"
    return value


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def ensure_sources(source_dir: Path) -> dict[str, Path]:
    source_dir.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for source in SOURCES:
        path = source_dir / source["local_name"]
        if not path.exists():
            urllib.request.urlretrieve(source["url"], path)
        paths[source["id"]] = path
    return paths


def make_block(
    document_id: str,
    suffix: str,
    locator_type: str,
    locator_label: str,
    original_text: str,
    **extra: Any,
) -> dict[str, Any]:
    return {
        "id": f"{document_id}:{suffix}",
        "document_id": document_id,
        "locator_type": locator_type,
        "locator_label": locator_label,
        "page_number": extra.get("page_number"),
        "section_path": extra.get("section_path"),
        "paragraph_index": extra.get("paragraph_index"),
        "sheet_name": extra.get("sheet_name"),
        "cell_range": extra.get("cell_range"),
        "original_text": compact(original_text),
        "text_hash": hashlib.sha256(compact(original_text).encode("utf-8")).hexdigest(),
    }


def make_fact(
    fact_id: str,
    title: str,
    value: str,
    block_id: str,
    module_code: str,
    fact_key: str,
    category: str,
) -> tuple[dict[str, Any], dict[str, Any]]:
    fact = {
        "id": fact_id,
        "module_code": module_code,
        "fact_key": fact_key,
        "title": title,
        "fact_value": compact(value, 420),
        "category": category,
        "status": "test_verified",
        "confidence": 100,
    }
    link = {
        "id": f"link:{fact_id}",
        "fact_id": fact_id,
        "source_block_id": block_id,
        "relation_type": "supports",
        "is_primary": 1,
    }
    return fact, link


def parse_pdf(source: dict[str, Any], path: Path):
    blocks: list[dict[str, Any]] = []
    facts: list[dict[str, Any]] = []
    links: list[dict[str, Any]] = []
    fact_pages = {
        10: ("装修辅材明细表", "辅材与常用品牌", "I04", "auxiliary-materials"),
        35: ("水电工程避坑", "水电施工经验", "I04", "water-electricity"),
        50: ("发泡水泥回填", "卫生间回填经验", "I04", "bathroom-backfill"),
        100: ("油漆工工程验收表", "油漆验收标准", "I06", "painting-acceptance"),
    }
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = compact(page.extract_text() or "")
            if not text:
                continue
            block = make_block(
                source["id"],
                f"page-{page_number:03d}",
                "page",
                f"第 {page_number} 页",
                text,
                page_number=page_number,
                section_path=f"第 {page_number} 页",
            )
            blocks.append(block)
            if page_number in fact_pages:
                title, category, module_code, fact_key = fact_pages[page_number]
                fact, link = make_fact(
                    f"fact:pdf:{page_number:03d}",
                    title,
                    text,
                    block["id"],
                    module_code,
                    fact_key,
                    category,
                )
                facts.append(fact)
                links.append(link)
    return blocks, facts, links, {"page_count": 100, "sheet_count": None}


SECTION_RE = re.compile(r"^[一二三四五六七八九十]+、")


def parse_docx(source: dict[str, Any], path: Path):
    document = Document(path)
    blocks: list[dict[str, Any]] = []
    facts: list[dict[str, Any]] = []
    links: list[dict[str, Any]] = []
    section = "正文"
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = compact(paragraph.text)
        if not text:
            continue
        if SECTION_RE.match(text):
            section = text
        block = make_block(
            source["id"],
            f"paragraph-{index:03d}",
            "paragraph",
            f"{section}｜第 {index} 段",
            text,
            section_path=section,
            paragraph_index=index,
        )
        blocks.append(block)
        if index >= 18 and (text[0].isdigit() or paragraph.style.name == "Body Text"):
            category = section.replace("、", " ", 1)
            fact, link = make_fact(
                f"fact:docx:{index:03d}",
                compact(text, 42),
                text,
                block["id"],
                "I06" if "验收" in section or "工程" in section else "I02",
                f"docx-paragraph-{index}",
                category,
            )
            facts.append(fact)
            links.append(link)

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [compact(cell.text) for cell in row.cells]
            text = "｜".join(value for value in values if value)
            if not text:
                continue
            block = make_block(
                source["id"],
                f"table-{table_index:02d}-row-{row_index:03d}",
                "table_row",
                f"表格 {table_index}｜第 {row_index} 行",
                text,
                section_path=f"表格 {table_index}",
            )
            blocks.append(block)
    return blocks, facts, links, {"page_count": None, "sheet_count": None}


def parse_xlsx(source: dict[str, Any], path: Path):
    workbook = load_workbook(path, data_only=False)
    blocks: list[dict[str, Any]] = []
    facts: list[dict[str, Any]] = []
    links: list[dict[str, Any]] = []
    for sheet in workbook.worksheets:
        current_room = ""
        current_part = ""
        for row_index in range(1, sheet.max_row + 1):
            values = [sheet.cell(row_index, column).value for column in range(1, sheet.max_column + 1)]
            display = [compact(str(value)) if value is not None else "" for value in values]
            if not any(display):
                continue
            current_room = display[1] or current_room
            if display[0] and display[1]:
                current_part = display[1]
            elif display[1]:
                current_part = display[1]
            start_col = next(index for index, value in enumerate(display, start=1) if value)
            end_col = max(index for index, value in enumerate(display, start=1) if value)
            cell_range = f"{get_column_letter(start_col)}{row_index}:{get_column_letter(end_col)}{row_index}"
            text = "｜".join(value for value in display if value)
            block = make_block(
                source["id"],
                f"{sheet.title}-row-{row_index:03d}",
                "cell_range",
                f"{sheet.title}!{cell_range}",
                text,
                sheet_name=sheet.title,
                cell_range=cell_range,
                section_path=current_room or current_part or sheet.title,
            )
            blocks.append(block)
            if row_index >= 5 and len(display) >= 4 and display[2] and display[3]:
                title_parts = [value for value in [current_room, current_part, display[2]] if value]
                title = " / ".join(dict.fromkeys(title_parts))
                fact, link = make_fact(
                    f"fact:xlsx:{sheet.title}:{row_index:03d}",
                    title,
                    f"{title}：{display[3]}",
                    block["id"],
                    "I06",
                    f"acceptance-row-{row_index}",
                    "精装修验收标准",
                )
                facts.append(fact)
                links.append(link)
    return blocks, facts, links, {"page_count": None, "sheet_count": len(workbook.sheetnames)}


def build_fixture(source_dir: Path) -> dict[str, Any]:
    paths = ensure_sources(source_dir)
    documents: list[dict[str, Any]] = []
    blocks: list[dict[str, Any]] = []
    facts: list[dict[str, Any]] = []
    links: list[dict[str, Any]] = []
    parsers = {"pdf": parse_pdf, "docx": parse_docx, "xlsx": parse_xlsx}

    for source in SOURCES:
        path = paths[source["id"]]
        source_blocks, source_facts, source_links, counts = parsers[source["file_type"]](source, path)
        blocks.extend(source_blocks)
        facts.extend(source_facts)
        links.extend(source_links)
        documents.append(
            {
                "id": source["id"],
                "title": source["title"],
                "original_filename": source["filename"],
                "file_url": source["url"],
                "file_type": source["file_type"],
                "category": source["category"],
                "sha256": sha256(path),
                "source_updated_at": source["updated_at"],
                "page_count": counts["page_count"],
                "sheet_count": counts["sheet_count"],
                "block_count": len(source_blocks),
                "fact_count": len(source_facts),
                "import_status": "verified_original",
            }
        )

    return {
        "fixture_version": "2026-08-02-v1",
        "scope": "isolated-provenance-test",
        "documents": documents,
        "blocks": blocks,
        "facts": facts,
        "evidence_links": links,
    }


def build_ui_fixture(fixture: dict[str, Any]) -> dict[str, Any]:
    documents = {item["id"]: item for item in fixture["documents"]}
    blocks = {item["id"]: item for item in fixture["blocks"]}
    links = {item["fact_id"]: item for item in fixture["evidence_links"]}
    grouped: dict[str, list[dict[str, Any]]] = {"pdf": [], "docx": [], "xlsx": []}
    for fact in fixture["facts"]:
        link = links[fact["id"]]
        block = blocks[link["source_block_id"]]
        document = documents[block["document_id"]]
        grouped[document["file_type"]].append(
            {
                **fact,
                "document": document,
                "evidence": block,
            }
        )
    records = grouped["pdf"][:4] + grouped["docx"][:12] + grouped["xlsx"][:20]
    return {
        "fixture_version": fixture["fixture_version"],
        "scope": fixture["scope"],
        "summary": {
            "documents": len(fixture["documents"]),
            "blocks": len(fixture["blocks"]),
            "facts": len(fixture["facts"]),
            "complete_evidence": len(fixture["evidence_links"]),
        },
        "documents": fixture["documents"],
        "records": records,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", type=Path, default=Path("/tmp/akke-source-samples"))
    parser.add_argument("--output", type=Path, default=Path("data/provenance-test-fixture.json"))
    parser.add_argument("--ui-output", type=Path, default=Path("data/provenance-test-ui.json"))
    args = parser.parse_args()
    fixture = build_fixture(args.source_dir)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(fixture, ensure_ascii=False, indent=2), encoding="utf-8")
    ui_fixture = build_ui_fixture(fixture)
    args.ui_output.parent.mkdir(parents=True, exist_ok=True)
    args.ui_output.write_text(json.dumps(ui_fixture, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {
                "documents": len(fixture["documents"]),
                "blocks": len(fixture["blocks"]),
                "facts": len(fixture["facts"]),
                "links": len(fixture["evidence_links"]),
                "output": str(args.output),
                "ui_output": str(args.ui_output),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
